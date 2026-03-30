"""
从多份 Word 文档中提取数据：段落、标题、表格。
提取结果写入 data/extracted/，每份 docx 对应一个 .md 文件，供后续根据模板做论文使用。
"""
from pathlib import Path

from docx import Document
from docx.document import Document as Doc

_ROOT = Path(__file__).resolve().parent.parent
INPUT_DOCX_DIR = _ROOT / "data" / "input"
EXTRACTED_DIR = _ROOT / "data" / "extracted"


def _table_to_markdown(tbl) -> str:
    rows = []
    for i, row in enumerate(tbl.rows):
        cells = [c.text.replace("|", "\\|").strip() for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n".join(rows)


def extract_one_docx(docx_path: Path) -> str:
    """从单份 Word 提取内容（段落+表格），返回 Markdown。"""
    doc: Doc = Document(str(docx_path))
    out: list[str] = []
    out.append("# " + docx_path.stem + "\n\n")

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = (p.style.name if p.style else None) or ""
        if "heading" in style_name.lower() or "标题" in style_name:
            level = 1
            for i in range(1, 10):
                if f"Heading {i}" in style_name or f"标题{i}" in style_name or f"Heading{i}" in style_name:
                    level = i
                    break
            out.append("\n" + "#" * min(level, 6) + " " + text + "\n")
        else:
            out.append(text + "\n\n")

    for tbl in doc.tables:
        out.append("\n" + _table_to_markdown(tbl) + "\n\n")

    return "".join(out).strip() or "（无正文）"


def read_docx_to_text(docx_path: Path | str) -> str:
    """将已生成的论文 Word 读成纯文本，供续写/纠偏时参考。"""
    doc = Document(str(docx_path))
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts)


def list_input_docx(input_dir: Path | None = None) -> list[Path]:
    """列出目录下所有 .docx 文件（不含以 ~ 开头）。不传则用默认 data/input。"""
    dir_path = input_dir if input_dir is not None else INPUT_DOCX_DIR
    if not dir_path.exists():
        return []
    return sorted(
        p for p in dir_path.glob("*.docx")
        if not p.name.startswith("~")
    )


def list_input_md_sql(input_dir: Path | None = None) -> list[Path]:
    """列出目录下所有 .md 和 .sql 文件（不含以 ~ 开头）。"""
    dir_path = input_dir if input_dir is not None else INPUT_DOCX_DIR
    if not dir_path.exists():
        return []
    return sorted(
        p for p in list(dir_path.glob("*.md")) + list(dir_path.glob("*.sql"))
        if not p.name.startswith("~")
    )


def sync_md_sql_to_output(input_dir: Path, output_dir: Path) -> list[Path]:
    """将 input_dir 下所有 .md 和 .sql 复制到 output_dir，返回复制后的路径列表。"""
    import shutil
    output_dir.mkdir(parents=True, exist_ok=True)
    copied: list[Path] = []
    for p in list_input_md_sql(input_dir):
        dest = output_dir / p.name
        shutil.copy2(p, dest)
        copied.append(dest)
    return copied


def run_extract(
    output_dir: Path | None = None,
    input_dir: Path | None = None,
) -> list[Path]:
    """
    从指定目录读取所有 Word，提取内容并写入 output_dir。
    不传 input_dir 则用 data/input；不传 output_dir 则用 data/extracted（若传了 input_dir 则用 data/extracted/input_dir.name）。
    """
    in_dir = input_dir if input_dir is not None else INPUT_DOCX_DIR
    out_dir = output_dir if output_dir is not None else (EXTRACTED_DIR / in_dir.name if input_dir is not None else EXTRACTED_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    generated: list[Path] = []
    for docx_path in list_input_docx(in_dir):
        try:
            content = extract_one_docx(docx_path)
            # print(f"word{docx_path.name}文字长度：{len(content)}")
            md_path = out_dir / (docx_path.stem + ".md")
            md_path.write_text(content, encoding="utf-8")
            generated.append(md_path)
        except Exception as e:
            raise RuntimeError(f"提取失败 {docx_path}: {e}") from e
    generated.extend(sync_md_sql_to_output(in_dir, out_dir))
    return generated
