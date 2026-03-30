"""
Word 文档构建器：基于 python-docx 封装常用 Word 文档操作。
支持：字体、字号、居中、两端对齐、首行缩进、行距、段前段后等。

对齐方式：使用 docx.enum.text.WD_ALIGN_PARAGRAPH
  - WD_ALIGN_PARAGRAPH.LEFT  左对齐
  - WD_ALIGN_PARAGRAPH.CENTER 居中
  - WD_ALIGN_PARAGRAPH.RIGHT 右对齐
  - WD_ALIGN_PARAGRAPH.JUSTIFY 两端对齐

字号常量：FONT_SIZE_16/15/14/12/10_5；字体常量：FONT_SONG、FONT_HEI、FONT_TIMES_NEW_ROMAN。
"""
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as Doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLACK = RGBColor(0, 0, 0)

# 将文本中的换行全部拆成「独立段落」，避免 python-docx 把 \n 写成段内软回车（Shift+Enter）
_LINE_SPLIT_RE = re.compile(r"[\r\n\u2028]+")


def _split_into_hard_paragraphs(text: str) -> list[str]:
    if text is None:
        return [""]
    if not _LINE_SPLIT_RE.search(text):
        return [text]
    parts = [p.strip() for p in _LINE_SPLIT_RE.split(text)]
    parts = [p for p in parts if p]
    return parts if parts else [""]


# 常用字号（磅 Pt），对应论文要求：三号/小三/四号/小四/五号
FONT_SIZE_16 = 16   # 三号
FONT_SIZE_15 = 15   # 小三
FONT_SIZE_14 = 14   # 四号
FONT_SIZE_12 = 12   # 小四
FONT_SIZE_10_5 = 10.5  # 五号

# 常用中英文字体名
FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_TIMES_NEW_ROMAN = "Times New Roman"

# 段落默认：行距 23 磅；段前 0 行、段后 0 行（行距，直接写 w:beforeLines/w:afterLines）；首行缩进 2 字符（w:firstLineChars）
DEFAULT_LINE_SPACING_PT = 23
DEFAULT_SPACE_BEFORE_LINES = 0 
DEFAULT_SPACE_AFTER_LINES = 0 
DEFAULT_FIRST_LINE_INDENT_CHARS = 2  # 首行缩进 2 字符（写入 w:firstLineChars，百分之一字符）


def _set_paragraph_runs_black(paragraph) -> None:
    """将段落内所有 run 的字体设为黑色。"""
    for run in paragraph.runs:
        run.font.color.rgb = BLACK


def _set_run_font(run, *, font_name_cn=None, font_name_en=None, font_size_pt=None):
    """设置 run 的字体、字号、颜色。

    Args:
        run: run 对象
        font_name_cn: 中文字体
        font_name_en: 英文字体
        font_size_pt: 字号
    """
    run.font.color.rgb = BLACK

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()

    if font_name_en:
        rFonts.set(qn("w:ascii"), font_name_en)
        rFonts.set(qn("w:hAnsi"), font_name_en)

    if font_name_cn:
        rFonts.set(qn("w:eastAsia"), font_name_cn)

    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def _set_first_line_indent_chars(paragraph, chars: float) -> None:
    """用字符设置首行缩进：直接写 OOXML 的 w:firstLineChars（百分之一字符），并清除 w:firstLine，让 Word 按字符显示。"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    ind = pPr.get_or_add_ind()
    # w:firstLineChars 单位是百分之一字符，2 字符 = 200
    ind.set(qn("w:firstLineChars"), str(int(round(chars * 100))))
    # 去掉 w:firstLine / w:hanging，否则 Word 可能用 twips 而不用字符
    for attr in ("w:firstLine", "w:hanging"):
        key = qn(attr)
        if ind.get(key) is not None:
            del ind.attrib[key]


_SINGLE_LINE_TWIPS = 240  # Word 标准单行高 12pt = 240 twips


def _set_spacing_before_after_lines(
    paragraph,
    before_lines: float | None = None,
    after_lines: float | None = None,
) -> None:
    """段前段后按行数设置。

    写入 w:beforeLines / w:afterLines（百分之一行）让 Word 显示"行"单位，
    同时写入 w:before / w:after（twips，按 240 twips/行）作为回退。
    """
    if before_lines is None and after_lines is None:
        return
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if before_lines is not None:
        spacing.set(qn("w:beforeLines"), str(int(round(before_lines * 100))))
        spacing.set(qn("w:before"), str(int(round(before_lines * _SINGLE_LINE_TWIPS))))
    if after_lines is not None:
        spacing.set(qn("w:afterLines"), str(int(round(after_lines * 100))))
        spacing.set(qn("w:after"), str(int(round(after_lines * _SINGLE_LINE_TWIPS))))



def _apply_paragraph_format(
    paragraph,
    *,
    alignment: int | None = None,
    line_spacing_pt: float | None = None,
    first_line_indent_chars: float | None = None,
    space_before_lines: float | None = None,
    space_after_lines: float | None = None,
) -> None:
    """设置段落格式：对齐、行距、首行缩进(字符)、段前段后(行)。"""
    pf = paragraph.paragraph_format
    if alignment is not None:
        paragraph.alignment = alignment
    if line_spacing_pt is not None:
        pf.line_spacing = Pt(line_spacing_pt)
    if first_line_indent_chars is not None:
        _set_first_line_indent_chars(paragraph, first_line_indent_chars)
    if space_before_lines is not None or space_after_lines is not None:
        _set_spacing_before_after_lines(
            paragraph,
            before_lines=space_before_lines,
            after_lines=space_after_lines,
        )


def _set_cell_border_top_bottom(cell, top: float | None = None, bottom: float | None = None) -> None:
    """
    为单元格设置上/下边框（用于三线表）。

    top / bottom 使用“线宽（磅）”，内部转换为 Word 的 1/8 pt 单位。
    """
    if top is None and bottom is None:
        return
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)

    def _add_edge(tag: str, size_pt: float) -> None:
        e = OxmlElement(tag)
        e.set(qn("w:val"), "single")
        # Word 边框单位：1/8 pt
        e.set(qn("w:sz"), str(int(round(size_pt * 8))))
        e.set(qn("w:color"), "000000")
        borders.append(e)

    if top is not None:
        _add_edge("w:top", top)
    if bottom is not None:
        _add_edge("w:bottom", bottom)


class WordBuilder:
    """Word 文档构建器，封装 python-docx 常用操作。"""

    def __init__(self) -> None:
        self._document: Doc = Document()

    @property
    def document(self) -> Doc:
        """获取底层 Document 对象，用于高级自定义操作。"""
        return self._document

    def add_heading(
        self,
        text: str,
        level: int = 0,
        *,
        font_name_cn: str | None = None,
        font_name_en: str | None = FONT_TIMES_NEW_ROMAN,
        font_size_pt: float | None = None,
        alignment: int | None = None,
    ) -> "WordBuilder":
        """
        添加标题。level=0 为章（默认居中），1 为节、2 为小节（默认顶格左对齐）。
        传 alignment 可覆盖默认。
        """
        style = self._document.styles['Title']
        pPr = style._element.get_or_add_pPr()

        for e in pPr.findall(".//w:pBdr", pPr.nsmap):
            pPr.remove(e)
        
        p = self._document.add_heading(text, level)
        # p.style = "Normal"
        # 未传字号时按 level：章 三号 16pt / 节 小三 15pt / 小节 四号 14pt
        if font_size_pt is None:
            font_size_pt = {0: FONT_SIZE_16, 1: FONT_SIZE_15, 2: FONT_SIZE_14}.get(level, FONT_SIZE_12)
        # 未传对齐时：章居中，节/小节顶格（左对齐）
        if alignment is None:
            alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.bold = False
            _set_run_font(
                run,
                font_name_cn=font_name_cn or FONT_HEI,
                font_name_en=font_name_en,
                font_size_pt=font_size_pt,
            )
        _apply_paragraph_format(p, alignment=alignment)
        return self

    def add_paragraph(
        self,
        text: str = "",
        style: str | None = None,
        *,
        font_name_cn: str | None = FONT_SONG,
        font_name_en: str | None = FONT_TIMES_NEW_ROMAN,
        font_size_pt: float | None = FONT_SIZE_12,
        alignment: int | None = WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent_chars: float | None = None,
        line_spacing_pt: float | None = DEFAULT_LINE_SPACING_PT,
        space_before_lines: float | None = None,
        space_after_lines: float | None = None,
    ) -> "WordBuilder":
        """
        添加普通段落。默认：段前 0.5 行、段后 1.5 行、首行缩进 2 字符、行距 23 磅。
        first_line_indent_chars：首行缩进字符数；space_before_lines/space_after_lines：段前/段后行数，None 用默认。
        """
        if first_line_indent_chars is None:
            first_line_indent_chars = DEFAULT_FIRST_LINE_INDENT_CHARS
        if space_before_lines is None:
            space_before_lines = DEFAULT_SPACE_BEFORE_LINES
        if space_after_lines is None:
            space_after_lines = DEFAULT_SPACE_AFTER_LINES

        pieces = _split_into_hard_paragraphs(text)
        for piece in pieces:
            if style:
                self._document.add_paragraph(piece, style=style)
            else:
                self._document.add_paragraph(piece)
            if not self._document.paragraphs:
                continue
            p = self._document.paragraphs[-1]
            for run in p.runs:
                _set_run_font(
                    run,
                    font_name_cn=font_name_cn,
                    font_name_en=font_name_en,
                    font_size_pt=font_size_pt,
                )
            _apply_paragraph_format(
                p,
                alignment=alignment,
                line_spacing_pt=line_spacing_pt,
                first_line_indent_chars=first_line_indent_chars,
                space_before_lines=space_before_lines,
                space_after_lines=space_after_lines,
            )
        return self

    def add_paragraph_with_runs(
        self,
        parts: list[tuple[str, bool, float | None, bool]],
        *,
        font_name_cn: str | None = FONT_SONG,
        font_name_en: str | None = FONT_TIMES_NEW_ROMAN,
        default_font_size_pt: float = FONT_SIZE_12,
        alignment: int | None = WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing_pt: float | None = 23,
    ) -> "WordBuilder":
        """
        添加由多段 run 组成的段落。
        parts: [(文本, 是否加粗, 字号Pt或None, 是否下划线), ...]，None 时用 default_font_size_pt。
        """
        p = self._document.add_paragraph()
        for text, bold, font_size_pt, underline in parts:
            run = p.add_run(text)
            run.bold = bold
            run.underline = underline
            _set_run_font(
                run,
                font_name_cn=font_name_cn,
                font_name_en=font_name_en,
                font_size_pt=font_size_pt if font_size_pt is not None else default_font_size_pt,
            )
        _apply_paragraph_format(p, alignment=alignment, line_spacing_pt=line_spacing_pt)
        return self

    def add_quote(
        self,
        text: str,
        *,
        font_name_cn: str | None = FONT_SONG,
        font_name_en: str | None = FONT_TIMES_NEW_ROMAN,
        font_size_pt: float | None = FONT_SIZE_12,
        alignment: int | None = None,
    ) -> "WordBuilder":
        """添加引用样式段落。可指定字体、字号、对齐。"""
        self._document.add_paragraph(text, style="Intense Quote")
        if self._document.paragraphs:
            p = self._document.paragraphs[-1]
            for run in p.runs:
                _set_run_font(run, font_name_cn=font_name_cn, font_name_en=font_name_en, font_size_pt=font_size_pt)
            if alignment is not None:
                _apply_paragraph_format(p, alignment=alignment)
        return self

    def add_bullet(
        self,
        text: str,
        *,
        font_name_cn: str | None = FONT_SONG,
        font_name_en: str | None = FONT_TIMES_NEW_ROMAN,
        font_size_pt: float | None = FONT_SIZE_12,
    ) -> "WordBuilder":
        """添加无序列表项。可指定字体、字号。"""
        self._document.add_paragraph(text, style="List Bullet")
        if self._document.paragraphs:
            p = self._document.paragraphs[-1]
            for run in p.runs:
                _set_run_font(run, font_name_cn=font_name_cn, font_name_en=font_name_en, font_size_pt=font_size_pt)
        return self

    def add_numbered(
        self,
        text: str,
        *,
        font_name_cn: str | None = FONT_SONG,
        font_name_en: str | None = FONT_TIMES_NEW_ROMAN,
        font_size_pt: float | None = FONT_SIZE_12,
    ) -> "WordBuilder":
        """添加有序列表项。可指定字体、字号。"""
        self._document.add_paragraph(text, style="List Number")
        if self._document.paragraphs:
            p = self._document.paragraphs[-1]
            for run in p.runs:
                _set_run_font(run, font_name_cn=font_name_cn, font_name_en=font_name_en, font_size_pt=font_size_pt)
        return self

    def add_picture(self, image_path: str | Path, width_cm: float = 5.2) -> "WordBuilder":
        """添加图片，width_cm 为宽度（厘米）。"""
        self._document.add_picture(str(image_path), width=Cm(width_cm))
        return self

    def add_section(self) -> "WordBuilder":
        """添加分节符。"""
        self._document.add_section()
        return self

    def add_table(
        self,
        headers: tuple[str, ...],
        rows: tuple[tuple[Any, ...], ...],
    ) -> "WordBuilder":
        """
        添加表格（默认使用符合论文排版的三线表样式）。

        headers: 表头元组，如 ('姓名', '性别', '出生日期')
        rows: 数据行元组的元组，如 (('骆昊', '男', '1995-5-5'), ...)
        """
        self.add_three_line_table(headers, rows)
        return self

    def add_three_line_table(
        self,
        headers: tuple[str, ...],
        rows: tuple[tuple[Any, ...], ...],
    ) -> "WordBuilder":
        """
        添加三线表：
        - 顶部粗线 + 表头下细线 + 底部粗线
        - 表头居中，数据行左对齐
        - 字体：中文宋体，英文 Times New Roman，字号 10.5 磅（五号）
        """
        cols = len(headers)
        total_rows = 1 + len(rows)
        table = self._document.add_table(rows=total_rows, cols=cols)

        # 设置表格宽度为页面宽度的 100% 左右（使用 5000 / 5000 = 100%）
        tbl = table._element
        tblPr = tbl.tblPr
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:type"), "pct")
        tblW.set(qn("w:w"), "5000")
        tblPr.append(tblW)

        # 填充数据并统一字体、对齐和行距
        data: list[list[Any]] = [list(headers)] + [list(r) for r in rows]
        for i in range(total_rows):
            for j in range(cols):
                cell = table.cell(i, j)
                # 清空默认段落内容，再写入
                p = cell.paragraphs[0]
                if p.runs:
                    # 清理现有 runs
                    for _ in range(len(p.runs)):
                        p.runs[0].clear()
                run = p.add_run(str(data[i][j]))
                _set_run_font(
                    run,
                    font_name_cn=FONT_SONG,
                    font_name_en=FONT_TIMES_NEW_ROMAN,
                    font_size_pt=FONT_SIZE_10_5,
                )

                # 表头水平居中，数据左对齐
                if i == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 行距等段落格式：单倍行距附近，段前/后为 0
                _apply_paragraph_format(
                    p,
                    alignment=p.alignment,
                    line_spacing_pt=18,
                    space_before_lines=0,
                    space_after_lines=0,
                )

        # 设置三线表边框：顶线 / 表头底线 / 最后一行底线
        top_bottom_pt = 1.5  # 顶线、底线：1.5 磅
        middle_pt = 0.5      # 表头线：0.5 磅

        # 顶线
        for c in table.rows[0].cells:
            _set_cell_border_top_bottom(c, top=top_bottom_pt)

        # 表头下细线
        for c in table.rows[0].cells:
            _set_cell_border_top_bottom(c, bottom=middle_pt)

        # 底线
        for c in table.rows[-1].cells:
            _set_cell_border_top_bottom(c, bottom=top_bottom_pt)

        return self

    def add_page_break(self) -> "WordBuilder":
        """添加分页符。"""
        self._document.add_page_break()
        return self

    def format_last_paragraph(
        self,
        *,
        alignment: int | None = None,
        line_spacing_pt: float | None = None,
        first_line_indent_chars: float | None = None,
        space_before_lines: float | None = None,
        space_after_lines: float | None = None,
    ) -> "WordBuilder":
        """对刚添加的最后一个段落设置格式。首行缩进用字符，段前段后用行，None 表示不改。"""
        if not self._document.paragraphs:
            return self
        p = self._document.paragraphs[-1]
        _apply_paragraph_format(
            p,
            alignment=alignment,
            line_spacing_pt=line_spacing_pt,
            first_line_indent_chars=first_line_indent_chars,
            space_before_lines=space_before_lines,
            space_after_lines=space_after_lines,
        )
        return self

    def save(self, path: str | Path) -> "WordBuilder":
        """保存文档到指定路径。"""
        self._document.save(str(path))
        return self
