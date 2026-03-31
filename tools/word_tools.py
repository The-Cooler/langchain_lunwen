"""
Word 相关的通用工具封装。

说明：
- 这里不保存对话状态，只负责提供操作 docx 的基础能力；
- 为 LangChain Agent 提供 WordBuilder 初始化、写节、写表等工具。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from tools.word_builder import WordBuilder

from tools.thesis_tools import append_progress, strip_thinking_from_content


def ensure_word_builder(
    current: Optional[WordBuilder],
    docx_path: str | Path,
    existing_docx_text: Optional[str],
) -> Tuple[WordBuilder, Path]:
    """
    确保返回一个可用的 WordBuilder 实例，并返回标准化后的 docx 路径。

    - 若 current 为 None，则创建新的 WordBuilder；
    - 若 docx_path 已存在且传入了 existing_docx_text，则从已有文档加载，避免覆盖。
    """
    path = Path(docx_path)
    builder = current or WordBuilder()
    if path.exists() and existing_docx_text:
        from docx import Document

        builder._document = Document(str(path))
    return builder, path


def _section_heading_level(section_title: str) -> int:
    """与论文规范一致：章→Heading 1，节→Heading 2，小节→Heading 3。

    注意：python-docx 的 add_heading(level=0) 会套用「标题/Title」样式，不是一级标题；
    因此章必须用 level=1 才对应 Word 的「标题 1」。
    """
    s = section_title.strip()
    if s.startswith("第") and "章" in s:
        return 1
    if re.match(r"^\d+\.\d+\.\d+", s):
        return 3
    if re.match(r"^\d+\.\d+", s):
        return 2
    # 摘要、关键词、致谢等：按一级标题处理
    return 1


def write_section_into_docx(
    builder: WordBuilder,
    docx_path: Path,
    section_title: str,
    content: str,
    progress_path: Optional[Path] = None,
) -> str:
    """
    将一节正式内容写入 Word。section_title 作为标题，content 为纯正文（可多段，用 \\n 换行）。
    严禁在 content 中写入思考、<thought> 等，仅限论文正文。
    """
    content = strip_thinking_from_content(content)
    if not section_title.strip() and not content:
        return "标题与内容为空，未写入。"
    level = _section_heading_level(section_title)
    builder.add_heading(section_title.strip() or "未命名节", level=level)
    if content:
        # 正文统一交给 add_paragraph：内部将任意 \n 拆成多个真正段落，禁止段内软回车
        builder.add_paragraph(content.strip())
    builder.save(docx_path)
    if progress_path is not None:
        append_progress(progress_path, "section", section_title.strip())
    return f"已写入：{section_title[:30]}…，并保存到 {docx_path}。"


def write_table_into_docx(
    builder: WordBuilder,
    docx_path: Path,
    headers: list[str] | tuple[str, ...],
    rows: list[list[Any]] | tuple[tuple[Any, ...], ...],
    caption: Optional[str] = None,
    progress_path: Optional[Path] = None,
) -> str:
    """
    在当前 Word 文档中插入一个三线表（按论文规范）。表头有几列，每行就必须有几列；缺列用 "" 补齐。
    """
    if caption:
        builder.add_paragraph(
            caption,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_chars=0,
            line_spacing_pt=18,
            space_before_lines=0.3,
            space_after_lines=0,
        )
    headers_tuple = tuple(str(h) for h in headers)
    norm_rows: list[tuple[Any, ...]] = []
    for row in rows:
        r_list = list(row)
        if len(r_list) < len(headers_tuple):
            r_list.extend([""] * (len(headers_tuple) - len(r_list)))
        norm_rows.append(tuple(r_list[: len(headers_tuple)]))
    builder.add_three_line_table(headers_tuple, tuple(norm_rows))
    builder.save(docx_path)
    if progress_path is not None:
        append_progress(progress_path, "table", (caption or "").strip())
    return f"已插入三线表（{len(norm_rows)} 行，{len(headers_tuple)} 列）到 {docx_path}。"


def write_figure_caption_into_docx(
    builder: WordBuilder,
    docx_path: Path,
    caption: str,
    progress_path: Optional[Path] = None,
) -> str:
    """
    写入图片题注（如“图3-1 系统架构图”）。
    规范：居中、五号、单倍行距、首行不缩进。
    """
    cap = (caption or "").strip()
    if not cap:
        return "未写入：caption 为空。"
    builder.add_paragraph(
        cap,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_chars=0,
        line_spacing_pt=18,
        space_before_lines=0.3,
        space_after_lines=0,
        font_size_pt=10.5,
    )
    builder.save(docx_path)
    if progress_path is not None:
        append_progress(progress_path, "figure_caption", cap)
    return f"已写入图题：{cap[:30]}… 到 {docx_path}。"


def truncate_docx_from_heading(docx_path: str | Path, heading_title: str) -> bool:
    """
    截断 docx：删除从 `heading_title` 对应的段落开始（包含该段落）到文档末尾的所有内容。

    返回值：是否成功找到并截断。
    """
    from docx import Document

    p = Path(docx_path)
    if not p.exists():
        return False

    doc = Document(str(p))
    target_title = (heading_title or "").strip()
    if not target_title:
        return False

    target_el = None
    for para in doc.paragraphs:
        if (para.text or "").strip() == target_title:
            target_el = para._element
            break
    if target_el is None:
        return False

    body = doc.element.body
    children = list(body)
    deleting = False
    for child in children:
        if child is target_el:
            deleting = True
        if deleting:
            body.remove(child)

    doc.save(str(p))
    return True


def truncate_progress_from_section(progress_path: str | Path, section_title: str) -> bool:
    """
    截断 progress.txt：删除从首次出现 `section_title`（kind == section）开始（包含该行）到末尾的所有记录。
    """
    pp = Path(progress_path)
    if not pp.exists():
        return False

    target = (section_title or "").strip()
    if not target:
        return False

    lines = pp.read_text(encoding="utf-8").splitlines()
    kept: list[str] = []
    found = False
    for ln in lines:
        if not ln.strip() or "\t" not in ln:
            continue
        kind, title = ln.split("\t", 1)
        title = title.strip()
        if not found and kind == "section" and title == target:
            found = True
            break
        kept.append(ln.strip())

    if not found:
        return False

    pp.write_text("\n".join(kept) + ("\n" if kept else ""), encoding="utf-8")
    return True

